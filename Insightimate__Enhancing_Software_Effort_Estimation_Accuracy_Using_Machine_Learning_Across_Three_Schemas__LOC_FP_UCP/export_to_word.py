#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export highlighting guide to Word document
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_highlight(run, color="yellow"):
    """Add highlight to a run"""
    # Get the XML element
    rPr = run._r.get_or_add_rPr()
    # Create highlight element
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), color)
    rPr.append(highlight)

def create_highlight_guide():
    # Create document
    doc = Document()
    
    # Add title
    title = doc.add_heading('DANH SÁCH CÁC ĐOẠN CẦN BÔI VÀNG TRONG PAPER_V2/MAIN.TEX', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('Tổng hợp các thay đổi theo yêu cầu của Reviewer')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.italic = True
    
    doc.add_paragraph()
    
    # Summary section
    doc.add_heading('1. TỔNG QUAN', 1)
    summary = doc.add_paragraph()
    summary.add_run('• Tổng số dòng cần highlight: ').bold = True
    summary.add_run('~600+ dòng (47% của paper)\n')
    summary.add_run('• Số phần cần viết lại hoàn toàn: ').bold = True
    summary.add_run('3 phần\n')
    summary.add_run('• Số bảng/công thức mới: ').bold = True
    summary.add_run('6 bảng + 3 công thức\n')
    summary.add_run('• Số subsection mới: ').bold = True
    summary.add_run('12 subsections')
    
    doc.add_page_break()
    
    # Section 1: Complete rewrites
    doc.add_heading('2. CÁC PHẦN CẦN BÔI VÀNG TOÀN BỘ', 1)
    
    doc.add_heading('2.1. Introduction (Dòng 85-128)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('VIẾT LẠI HOÀN TOÀN - Thêm 3 objectives rõ ràng, thêm 3 methodological gaps, cấu trúc lại toàn bộ phần mở đầu')
    
    p = doc.add_paragraph()
    run = p.add_run('Nội dung cần bôi vàng:\n\n')
    run.font.size = Pt(11)
    
    intro_text = """However, three methodological gaps limit reproducibility in prior SEE research: (i) unclear dataset provenance and deduplication rules hinder independent replication; (ii) COCOMO~II baselines often use arbitrary default parameters when cost drivers are unavailable, creating unfair comparisons; (iii) cross-schema aggregation protocols (macro vs. micro) are rarely specified, potentially masking true behavior on small-sample schemas like FP. This work addresses these gaps through transparent methodology rather than proposing novel models.

The contributions of this paper are summarized as follows:
• We propose a unified multi-schema machine-learning framework that harmonizes preprocessing, feature construction, model training, and evaluation across LOC, FP, and UCP, aggregating n=3,054 projects from 18 sources (1993–2022) with explicit provenance and deduplication rules."""
    
    run = p.add_run(intro_text)
    add_highlight(run)
    run.font.size = Pt(10)
    
    # Section 2: Discussion
    doc.add_heading('2.2. Discussion (Dòng 1140-1210)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('VIẾT LẠI HOÀN TOÀN - Thêm phần Strengths/Weaknesses/Implications, cấu trúc lại hoàn toàn')
    
    p = doc.add_paragraph()
    run = p.add_run('Cần bôi vàng toàn bộ phần Discussion mới với cấu trúc:\n')
    run.font.size = Pt(11)
    p.add_run('• Strengths of the Proposed Framework\n')
    p.add_run('• Weaknesses and Limitations\n')
    p.add_run('• Implications for Practitioners')
    
    # Section 3: Data Availability
    doc.add_heading('2.3. Data Availability + Declarations (Dòng 1212-1286)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('PHẦN MỚI HOÀN TOÀN - Thêm Data Availability Statement, Funding, Conflicts of Interest, Author Contributions, Acknowledgments')
    
    p = doc.add_paragraph()
    run = p.add_run('Nội dung cần bôi vàng:\n\n')
    run.font.size = Pt(11)
    
    data_text = """All datasets, preprocessing scripts, trained models, and evaluation code are publicly available at:
• Dataset repository: https://github.com/Huy-VNNIC/AI-Project/tree/main/datasets
• Trained models: https://github.com/Huy-VNNIC/AI-Project/tree/main/models
• Preprocessing pipeline: https://github.com/Huy-VNNIC/AI-Project/tree/main/src
• Evaluation scripts: https://github.com/Huy-VNNIC/AI-Project/tree/main/scripts"""
    
    run = p.add_run(data_text)
    add_highlight(run)
    run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # Section 3: New subsections
    doc.add_heading('3. CÁC SUBSECTION/ĐOẠN MỚI CẦN BÔI VÀNG', 1)
    
    # 3.1 Baseline Calibration
    doc.add_heading('3.1. Baseline Fairness and Calibration (Dòng 133-143)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Subsection MỚI - Giải thích cách calibrate baseline')
    
    p = doc.add_paragraph()
    run = p.add_run('Nội dung chính cần bôi vàng:\n\n')
    run.font.size = Pt(11)
    
    baseline_text = """To ensure fair comparison, we do not use full COCOMO~II with default parameters (most public datasets lack the required cost drivers and scale factors). Instead, we employ a calibrated size-only power-law baseline of the form E = A × (Size)^B, where A and B are fitted via least-squares regression on log(E) vs. log(Size) using training data only."""
    
    run = p.add_run(baseline_text)
    add_highlight(run)
    run.font.size = Pt(10)
    
    # 3.2 Comparative table
    doc.add_heading('3.2. Related Work Comparison Table (Dòng 145-178)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Bảng so sánh MỚI - So sánh với các nghiên cứu trước')
    p = doc.add_paragraph('Cần bôi vàng toàn bộ bảng comparative survey')
    
    # 3.3 Macro-averaging
    doc.add_heading('3.3. Cross-Schema Aggregation Protocol (Dòng 229-246)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Công thức + mô tả MỚI - Macro-averaging protocol')
    
    p = doc.add_paragraph()
    run = p.add_run('Công thức cần bôi vàng:\n\n')
    run.font.size = Pt(11)
    
    macro_text = """m_macro = (1/3) × Σ m^(s) for s ∈ {LOC, FP, UCP}

where m^(s) is the metric value for schema s. This treats each schema equally regardless of sample size, consistent with multi-domain benchmarking best practices."""
    
    run = p.add_run(macro_text)
    add_highlight(run)
    run.font.size = Pt(10)
    
    # 3.4 Dataset provenance table
    doc.add_heading('3.4. Dataset Provenance Manifest (Dòng 248-275)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Table 1 MỚI - Dataset summary với provenance')
    p = doc.add_paragraph('Cần bôi vàng toàn bộ Table 1 và caption')
    
    # 3.5 MdMRE formula
    doc.add_heading('3.5. MdMRE và MAPE Formulas (Trong phần 2.3)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('2 công thức MỚI - Equations cho MdMRE và MAPE')
    p = doc.add_paragraph()
    run = p.add_run('MdMRE = Median(|y_i - ŷ_i| / y_i)\nMAPE = (100%/n) × Σ |y_i - ŷ_i| / y_i')
    add_highlight(run)
    
    doc.add_page_break()
    
    # Section 4: Modified sections
    doc.add_heading('4. CÁC PHẦN SỬA ĐỔI/BỔ SUNG', 1)
    
    # 4.1 Results table
    doc.add_heading('4.1. Results Table - Thêm Cột và Dòng (Dòng 630-655)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Bổ sung vào bảng kết quả:\n')
    p.add_run('• Thêm cột MdMRE\n')
    p.add_run('• Thêm cột MAPE\n')
    p.add_run('• Thêm dòng XGBoost\n')
    p.add_run('• Xóa cột R²')
    
    p = doc.add_paragraph()
    run = p.add_run('Cần bôi vàng: Toàn bộ cột MdMRE, cột MAPE, và dòng XGBoost trong bảng kết quả')
    run.font.size = Pt(10)
    add_highlight(run)
    
    # 4.2 Ablation study
    doc.add_heading('4.2. Ablation Study + Feature Importance (Dòng 865-920)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Subsection MỚI - Ablation experiment và feature importance analysis')
    p = doc.add_paragraph('Cần bôi vàng toàn bộ subsection này bao gồm text mô tả + figure về feature importance')
    
    # 4.3 Dataset imbalance
    doc.add_heading('4.3. Dataset Imbalance Justification (Trong phần 2.3)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Đoạn văn MỚI - Giải thích về imbalance giữa LOC/FP/UCP')
    
    p = doc.add_paragraph()
    run = p.add_run('Đoạn văn cần bôi vàng bắt đầu với:\n\n')
    run.font.size = Pt(11)
    
    imbalance_text = """"The imbalance across schemas—LOC comprises 2,765 projects (90.5%), while FP contains 158 (5.2%) and UCP 131 (4.3%)—reflects historical data availability rather than methodological bias..."""
    
    run = p.add_run(imbalance_text)
    add_highlight(run)
    run.font.size = Pt(10)
    
    # 4.4 Confidence intervals
    doc.add_heading('4.4. Confidence Intervals Paragraph (Trong phần 2.3)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Đoạn văn MỚI - Mô tả về bootstrap CI')
    p = doc.add_paragraph()
    run = p.add_run('Paragraph về "All reported metrics include mean ± standard deviation across 10 random seeds..."')
    add_highlight(run)
    
    # 4.5 Leave-One-Source-Out
    doc.add_heading('4.5. Leave-One-Source-Out Validation (Nếu có trong Results)', 2)
    p = doc.add_paragraph()
    p.add_run('Mô tả: ').bold = True
    p.add_run('Subsection/paragraph MỚI - Cross-source validation')
    p = doc.add_paragraph('Tìm và bôi vàng phần mô tả LOSO validation với 11 LOC sources')
    
    doc.add_page_break()
    
    # Section 5: Quick reference table
    doc.add_heading('5. BẢNG TÓM TẮT NHANH', 1)
    
    # Create table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Dòng'
    hdr_cells[1].text = 'Phần'
    hdr_cells[2].text = 'Loại thay đổi'
    hdr_cells[3].text = 'Mức độ'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Data rows
    data = [
        ('70-84', 'Abstract', 'Sửa đổi nhỏ', 'Quan trọng'),
        ('85-128', 'Introduction', 'VIẾT LẠI HOÀN TOÀN', 'CỰC KỲ QUAN TRỌNG'),
        ('133-143', 'Baseline Calibration', 'Subsection MỚI', 'Quan trọng'),
        ('145-178', 'Related Work Table', 'Bảng MỚI', 'Quan trọng'),
        ('229-246', 'Macro-averaging', 'Công thức MỚI', 'Quan trọng'),
        ('248-275', 'Table 1 (Dataset)', 'Bảng MỚI', 'Rất quan trọng'),
        ('Trong 2.3', 'MdMRE/MAPE', '2 công thức MỚI', 'Quan trọng'),
        ('Trong 2.3', 'Imbalance paragraph', 'Đoạn văn MỚI', 'Quan trọng'),
        ('Trong 2.3', 'Confidence Intervals', 'Đoạn văn MỚI', 'Quan trọng'),
        ('630-655', 'Results Table', 'Thêm cột/dòng', 'Rất quan trọng'),
        ('865-920', 'Ablation Study', 'Subsection MỚI', 'Rất quan trọng'),
        ('1140-1210', 'Discussion', 'VIẾT LẠI HOÀN TOÀN', 'CỰC KỲ QUAN TRỌNG'),
        ('1212-1286', 'Data Availability', 'Phần MỚI HOÀN TOÀN', 'CỰC KỲ QUAN TRỌNG'),
    ]
    
    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            # Highlight important rows
            if 'HOÀN TOÀN' in text or 'CỰC KỲ' in text:
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(255, 0, 0)
    
    # Final notes
    doc.add_page_break()
    doc.add_heading('6. GHI CHÚ QUAN TRỌNG', 1)
    
    notes = doc.add_paragraph()
    notes.add_run('1. ').bold = True
    notes.add_run('Các phần đánh dấu "VIẾT LẠI HOÀN TOÀN" cần bôi vàng toàn bộ (100% nội dung)\n\n')
    
    notes.add_run('2. ').bold = True
    notes.add_run('Các bảng và công thức mới cần bôi vàng bao gồm cả caption và equation number\n\n')
    
    notes.add_run('3. ').bold = True
    notes.add_run('Tổng cộng khoảng 600+ dòng cần highlight, chiếm 47% paper\n\n')
    
    notes.add_run('4. ').bold = True
    notes.add_run('Sử dụng màu vàng (yellow) cho tất cả các phần được liệt kê\n\n')
    
    notes.add_run('5. ').bold = True
    notes.add_run('File gốc: Paper_v2/main.tex (1,286 dòng)\n\n')
    
    notes.add_run('6. ').bold = True
    notes.add_run('Tài liệu tham khảo: RESPONSE_REVIEWERS_COMPLETE.tex\n\n')
    
    # Contact info
    doc.add_paragraph()
    contact = doc.add_paragraph()
    contact.add_run('Tạo ngày: 19 Tháng 2, 2026\n').italic = True
    contact.add_run('Nguồn: Phân tích từ RESPONSE_REVIEWERS_COMPLETE.tex').italic = True
    
    # Save document
    output_path = 'HIGHLIGHT_GUIDE.docx'
    doc.save(output_path)
    print(f"✓ Đã tạo file {output_path} thành công!")
    print(f"✓ File chứa danh sách đầy đủ các phần cần bôi vàng")
    print(f"✓ Bao gồm: line numbers, mô tả, và nội dung text")
    
    return output_path

if __name__ == '__main__':
    create_highlight_guide()

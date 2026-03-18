#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Export Paper_v2/main.tex to Word document with highlighting
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def add_highlight(run, color="yellow"):
    """Add highlight to a run"""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement('w:highlight')
    highlight.set(qn('w:val'), color)
    rPr.append(highlight)

def should_highlight(line_num):
    """Check if a line should be highlighted based on reviewer changes"""
    highlight_ranges = [
        (70, 84),      # Abstract modifications
        (85, 128),     # Introduction REWRITE
        (133, 143),    # Baseline Calibration
        (145, 178),    # Related Work Table
        (229, 246),    # Macro-averaging formula
        (248, 275),    # Table 1 Dataset provenance
        (630, 655),    # Results table modifications
        (865, 920),    # Ablation study
        (1140, 1210),  # Discussion REWRITE
        (1212, 1286),  # Data Availability NEW
    ]
    
    for start, end in highlight_ranges:
        if start <= line_num <= end:
            return True
    return False

def clean_latex(text):
    """Remove common LaTeX commands for readable Word output"""
    # Remove comments
    text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
    
    # Remove common commands
    text = re.sub(r'\\cite\{[^}]+\}', '[ref]', text)
    text = re.sub(r'\\ref\{[^}]+\}', '[ref]', text)
    text = re.sub(r'\\label\{[^}]+\}', '', text)
    text = re.sub(r'\\textbf\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textit\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\emph\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\url\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\textsuperscript\{([^}]+)\}', r'^\1', text)
    text = re.sub(r'\\~', ' ', text)
    text = re.sub(r'~', ' ', text)
    
    # Clean whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    
    return text

def add_latex_paragraph(doc, text, line_num, is_heading=False, level=1):
    """Add a paragraph with appropriate formatting and highlighting"""
    text = clean_latex(text)
    if not text:
        return
    
    if is_heading:
        p = doc.add_heading(text, level=level)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(11)
    
    # Highlight if in changed section
    if should_highlight(line_num):
        for run in p.runs:
            add_highlight(run, "yellow")

def create_paper_word():
    # Read the LaTeX file
    tex_file = 'Paper_v2/main.tex'
    with open(tex_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)
    
    # Title
    title = doc.add_heading('Insightimate: Enhancing Software Effort Estimation Accuracy Using Machine Learning Across Three Schemas (LOC/FP/UCP)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(16)
        run.font.bold = True
    
    # Authors
    authors_p = doc.add_paragraph()
    authors_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors_run = authors_p.add_run(
        'Nguyen Nhat Huy¹, Duc Man Nguyen¹, Dang Nhat Minh¹, '
        'Nguyen Thuy Giang¹, P. W. C. Prasad¹, Md Shohel Sayeed²,*'
    )
    authors_run.font.size = Pt(12)
    
    # Affiliations
    affil_p = doc.add_paragraph()
    affil_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    affil_run = affil_p.add_run(
        '¹International School, Duy Tan University, Da Nang 550000, Vietnam\n'
        '²Faculty of Information Science and Technology, Multimedia University, Malaysia\n'
        '*Corresponding author: shohel.sayeed@mmu.edu.my'
    )
    affil_run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    # Process content
    current_section = None
    in_abstract = False
    in_equation = False
    in_figure = False
    in_table = False
    skip_lines = 0
    
    for i, line in enumerate(lines, 1):
        if skip_lines > 0:
            skip_lines -= 1
            continue
        
        line = line.strip()
        
        # Skip preamble
        if i < 60:
            continue
        
        # Skip end document
        if '\\end{document}' in line:
            break
        
        # Skip commands
        if line.startswith('%') or not line:
            continue
        
        # Track environments
        if '\\begin{abstract}' in line:
            in_abstract = True
            doc.add_heading('Abstract', 1)
            if should_highlight(i):
                for run in doc.paragraphs[-1].runs:
                    add_highlight(run)
            continue
        elif '\\end{abstract}' in line:
            in_abstract = False
            continue
        elif '\\begin{equation}' in line or '\\begin{align}' in line or '\\[' in line:
            in_equation = True
            continue
        elif '\\end{equation}' in line or '\\end{align}' in line or '\\]' in line:
            in_equation = False
            continue
        elif '\\begin{figure' in line:
            in_figure = True
            continue
        elif '\\end{figure}' in line:
            in_figure = False
            continue
        elif '\\begin{table' in line:
            in_table = True
            continue
        elif '\\end{table}' in line:
            in_table = False
            continue
        
        # Skip environments
        if in_equation or in_figure or in_table:
            continue
        
        # Section headings
        if '\\section{' in line:
            match = re.search(r'\\section\{([^}]+)\}', line)
            if match:
                heading = match.group(1)
                doc.add_heading(clean_latex(heading), 1)
                if should_highlight(i):
                    for run in doc.paragraphs[-1].runs:
                        add_highlight(run)
            continue
        elif '\\subsection{' in line:
            match = re.search(r'\\subsection\{([^}]+)\}', line)
            if match:
                heading = match.group(1)
                doc.add_heading(clean_latex(heading), 2)
                if should_highlight(i):
                    for run in doc.paragraphs[-1].runs:
                        add_highlight(run)
            continue
        elif '\\subsubsection{' in line:
            match = re.search(r'\\subsubsection\{([^}]+)\}', line)
            if match:
                heading = match.group(1)
                doc.add_heading(clean_latex(heading), 3)
                if should_highlight(i):
                    for run in doc.paragraphs[-1].runs:
                        add_highlight(run)
            continue
        elif '\\paragraph{' in line:
            match = re.search(r'\\paragraph\{([^}]+)\}', line)
            if match:
                heading = match.group(1)
                p = doc.add_paragraph()
                run = p.add_run(clean_latex(heading))
                run.font.bold = True
                run.font.size = Pt(11)
                if should_highlight(i):
                    add_highlight(run)
            continue
        
        # Skip LaTeX commands
        if line.startswith('\\'):
            continue
        
        # Regular content
        if line and not line.startswith('%'):
            # Accumulate text for paragraph
            paragraph_text = []
            start_line = i
            
            # Collect lines until empty line or special command
            while i <= len(lines):
                curr_line = lines[i-1].strip() if i <= len(lines) else ''
                
                if not curr_line or curr_line.startswith('\\section') or \
                   curr_line.startswith('\\subsection') or curr_line.startswith('\\begin') or \
                   curr_line.startswith('\\end') or curr_line.startswith('\\paragraph'):
                    break
                
                if not curr_line.startswith('%') and not curr_line.startswith('\\'):
                    paragraph_text.append(curr_line)
                
                i += 1
            
            skip_lines = i - start_line - 1
            
            if paragraph_text:
                full_text = ' '.join(paragraph_text)
                full_text = clean_latex(full_text)
                
                if full_text and len(full_text) > 10:
                    p = doc.add_paragraph()
                    run = p.add_run(full_text)
                    run.font.size = Pt(11)
                    
                    # Highlight if in changed range
                    if should_highlight(start_line):
                        add_highlight(run, "yellow")
    
    # Add note at end
    doc.add_page_break()
    note = doc.add_heading('GHI CHÚ', 1)
    note_p = doc.add_paragraph()
    note_run = note_p.add_run(
        'Các phần được tô vàng là những thay đổi đã thực hiện theo yêu cầu của Reviewer.\n\n'
        'Bao gồm:\n'
        '• Abstract (dòng 70-84): Sửa đổi nhỏ\n'
        '• Introduction (dòng 85-128): VIẾT LẠI HOÀN TOÀN\n'
        '• Baseline Calibration (dòng 133-143): Subsection mới\n'
        '• Related Work Table (dòng 145-178): Bảng so sánh mới\n'
        '• Macro-averaging Protocol (dòng 229-246): Công thức mới\n'
        '• Dataset Provenance Table (dòng 248-275): Table 1 mới\n'
        '• Results Table (dòng 630-655): Thêm cột MdMRE/MAPE và dòng XGBoost\n'
        '• Ablation Study (dòng 865-920): Subsection mới\n'
        '• Discussion (dòng 1140-1210): VIẾT LẠI HOÀN TOÀN\n'
        '• Data Availability (dòng 1212-1286): Phần mới hoàn toàn\n'
    )
    note_run.font.size = Pt(11)
    
    # Save document
    output_path = 'PAPER_V2_WITH_HIGHLIGHTS.docx'
    doc.save(output_path)
    print(f"✓ Đã tạo file {output_path} thành công!")
    print(f"✓ File chứa toàn bộ nội dung paper với các phần đã thay đổi được highlight màu vàng")
    
    return output_path

if __name__ == '__main__':
    create_paper_word()

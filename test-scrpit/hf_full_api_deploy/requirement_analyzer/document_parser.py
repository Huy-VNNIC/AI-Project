"""
Document parser module for handling different file formats in the requirement analyzer.
Supports .txt, .doc, .docx, .pdf, .md files.
"""

import io
import os
import tempfile
from pathlib import Path
import logging

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentParser:
    """
    Parser for extracting text from different document formats.
    """
    
    def __init__(self):
        """Initialize the document parser."""
        pass
        
    def parse(self, content, filename):
        """
        Parse document content based on file extension.
        
        Args:
            content (bytes): The file content as bytes
            filename (str): Original filename with extension
            
        Returns:
            str: Extracted text content
        
        Raises:
            ValueError: If file format is not supported
        """
        # Get file extension
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        
        # Route to appropriate parser based on extension
        if ext == '.txt' or ext == '.md':
            return self._parse_text(content)
        elif ext == '.pdf':
            return self._parse_pdf(content)
        elif ext == '.doc' or ext == '.docx':
            return self._parse_doc(content)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _parse_text(self, content):
        """Parse plain text or markdown files."""
        try:
            # Try UTF-8 first
            return content.decode('utf-8')
        except UnicodeDecodeError:
            # Fall back to Latin-1 if UTF-8 fails
            try:
                return content.decode('latin-1')
            except UnicodeDecodeError:
                logger.error("Failed to decode text with UTF-8 and Latin-1")
                raise ValueError("Unable to decode text file. The file might be corrupted or in an unsupported encoding.")
    
    def _parse_pdf(self, content):
        """Parse PDF files."""
        try:
            # Import PyPDF2 here to avoid dependency issues if not used
            import PyPDF2
            
            pdf_file = io.BytesIO(content)
            reader = PyPDF2.PdfReader(pdf_file)
            text = ""
            
            # Extract text from all pages
            for page_num in range(len(reader.pages)):
                text += reader.pages[page_num].extract_text() + "\n\n"
                
            return text
        except ImportError:
            logger.error("PyPDF2 library not installed. Cannot parse PDF.")
            raise ValueError("PDF parsing requires PyPDF2 library. Please install it using: pip install PyPDF2")
        except Exception as e:
            logger.error(f"Error parsing PDF: {str(e)}")
            raise ValueError(f"Failed to parse PDF file: {str(e)}")
    
    def _parse_doc(self, content):
        """Parse DOC/DOCX files."""
        try:
            # Import docx here to avoid dependency issues if not used
            import docx
            
            # Save content to a temporary file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp:
                temp.write(content)
                temp_path = temp.name
            
            # Parse the document
            try:
                doc = docx.Document(temp_path)
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
                            
                return text
            finally:
                # Clean up the temporary file
                if os.path.exists(temp_path):
                    os.unlink(temp_path)
                    
        except ImportError:
            logger.error("python-docx library not installed. Cannot parse DOC/DOCX.")
            raise ValueError("DOC/DOCX parsing requires python-docx library. Please install it using: pip install python-docx")
        except Exception as e:
            logger.error(f"Error parsing DOC/DOCX: {str(e)}")
            raise ValueError(f"Failed to parse DOC/DOCX file: {str(e)}")
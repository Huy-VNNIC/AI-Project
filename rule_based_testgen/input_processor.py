"""
Input Processing Module
Extract text from PDF, DOCX, or TXT files
"""

from typing import Union
from pathlib import Path


def extract_text_from_pdf(filepath: str) -> str:
    """
    Extract text from PDF file using pdfplumber
    
    Args:
        filepath: Path to PDF file
        
    Returns:
        Extracted text as string
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    
    text = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    text.append(f"--- Page {page_num} ---\n{page_text}")
    except Exception as e:
        raise ValueError(f"Error reading PDF {filepath}: {e}")
    
    return "\n".join(text)


def extract_text_from_docx(filepath: str) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        filepath: Path to DOCX file
        
    Returns:
        Extracted text as string
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    text = []
    try:
        doc = Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())
    except Exception as e:
        raise ValueError(f"Error reading DOCX {filepath}: {e}")
    
    return "\n".join(text)


def extract_text_from_txt(filepath: str) -> str:
    """
    Extract text from plain text file
    
    Args:
        filepath: Path to TXT file
        
    Returns:
        Extracted text as string
    """
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding
        with open(filepath, "r", encoding="latin-1") as f:
            return f.read()
    except Exception as e:
        raise ValueError(f"Error reading TXT {filepath}: {e}")


class InputProcessor:
    """Main input processing class"""
    
    SUPPORTED_FORMATS = {".pdf": extract_text_from_pdf, 
                         ".docx": extract_text_from_docx,
                         ".txt": extract_text_from_txt}
    
    @staticmethod
    def extract_text(filepath: str) -> str:
        """
        Extract text from file (auto-detect format)
        
        Args:
            filepath: Path to input file (PDF, DOCX, or TXT)
            
        Returns:
            Extracted raw text
            
        Raises:
            ValueError: If file format not supported
        """
        path = Path(filepath)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {filepath}")
        
        suffix = path.suffix.lower()
        
        if suffix not in InputProcessor.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {suffix}. "
                           f"Supported: {', '.join(InputProcessor.SUPPORTED_FORMATS.keys())}")
        
        extractor = InputProcessor.SUPPORTED_FORMATS[suffix]
        return extractor(filepath)
    
    @staticmethod
    def extract_from_bytes(data: bytes, format: str) -> str:
        """
        Extract text from bytes (for file uploads)
        
        Args:
            data: File content as bytes
            format: File format (pdf, docx, txt)
            
        Returns:
            Extracted text
        """
        import tempfile
        import os
        
        format = format.lower().lstrip(".")
        temp_path = tempfile.mktemp(suffix=f".{format}")
        
        try:
            with open(temp_path, "wb") as f:
                f.write(data)
            return InputProcessor.extract_text(temp_path)
        finally:
            if os.path.exists(temp_path):
                os.remove(temp_path)

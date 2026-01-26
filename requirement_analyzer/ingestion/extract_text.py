"""
File Text Extraction Module
Extract text from various file formats (txt, docx, pdf)
"""
import io
from typing import Optional
import logging

logger = logging.getLogger(__name__)


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    Extract text from .txt or .md files
    
    Args:
        file_bytes: File content as bytes
        
    Returns:
        Decoded text string
    """
    try:
        return file_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        logger.error(f"Error decoding text file: {e}")
        return ""


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text from .docx files
    
    Requires: python-docx
    
    Args:
        file_bytes: File content as bytes
        
    Returns:
        Extracted text string
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        return ""
    
    try:
        doc = Document(io.BytesIO(file_bytes))
        
        # Extract paragraphs
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        
        # Extract tables (optional - requirements might be in tables)
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)
        
        # Combine
        all_text = '\n'.join(paragraphs)
        if table_texts:
            all_text += '\n\n' + '\n'.join(table_texts)
        
        return all_text
    
    except Exception as e:
        logger.error(f"Error extracting from DOCX: {e}")
        return ""


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF files
    
    Tries multiple libraries in order:
    1. pymupdf (fitz) - fastest and most reliable
    2. pdfplumber - good for tables
    3. PyPDF2 - fallback
    
    Args:
        file_bytes: File content as bytes
        
    Returns:
        Extracted text string
    """
    # Try pymupdf (fitz) first
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text_parts = []
        
        for page in doc:
            text_parts.append(page.get_text())
        
        doc.close()
        return '\n\n'.join(text_parts)
    
    except ImportError:
        logger.debug("pymupdf not available, trying pdfplumber...")
    except Exception as e:
        logger.warning(f"pymupdf extraction failed: {e}")
    
    # Try pdfplumber
    try:
        import pdfplumber
        
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return '\n\n'.join(text_parts)
    
    except ImportError:
        logger.debug("pdfplumber not available, trying PyPDF2...")
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed: {e}")
    
    # Fallback to PyPDF2
    try:
        from PyPDF2 import PdfReader
        
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        return '\n\n'.join(text_parts)
    
    except ImportError:
        logger.error("No PDF library available. Install with: pip install pymupdf pdfplumber PyPDF2")
        return ""
    except Exception as e:
        logger.error(f"PyPDF2 extraction failed: {e}")
        return ""


def extract_text(filename: str, file_bytes: bytes) -> str:
    """
    Auto-detect file type and extract text
    
    Args:
        filename: Original filename (for extension detection)
        file_bytes: File content as bytes
        
    Returns:
        Extracted text string
    """
    filename_lower = filename.lower()
    
    if filename_lower.endswith('.txt') or filename_lower.endswith('.md'):
        return extract_text_from_txt(file_bytes)
    
    elif filename_lower.endswith('.docx'):
        return extract_text_from_docx(file_bytes)
    
    elif filename_lower.endswith('.pdf'):
        return extract_text_from_pdf(file_bytes)
    
    else:
        # Unknown extension, try as text
        logger.warning(f"Unknown file extension: {filename}. Trying as text...")
        return extract_text_from_txt(file_bytes)
